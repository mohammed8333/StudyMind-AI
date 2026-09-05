import io
import os
import pytest
import pytest_asyncio
from httpx import AsyncClient, ASGITransport
from PIL import Image, ImageDraw

from app.main import app
from app.core.database import init_db
from app.services.file_validator import sanitize_filename, validate_uploaded_file
from docx import Document as DocxDocument

@pytest_asyncio.fixture(autouse=True)
async def setup_db():
    await init_db()

async def get_or_create_user(ac: AsyncClient, email: str, name: str) -> str:
    await ac.post("/api/v1/auth/register", json={
        "email": email,
        "password": "Password123!",
        "full_name": name,
        "grade_or_level": "الثانوية العامة"
    })
    login_res = await ac.post("/api/v1/auth/login", data={
        "username": email,
        "password": "Password123!"
    })
    return login_res.json()["access_token"]

# -------------------------------------------------------------
# 1. Direct Unit Tests for File Validation & Path Sanitization
# -------------------------------------------------------------

def test_sanitize_filename_prevents_traversal():
    assert sanitize_filename("../../etc/passwd") == "passwd"
    assert sanitize_filename("..\\..\\windows\\system32.dll") == "system32.dll"
    assert sanitize_filename("safe_document.pdf") == "safe_document.pdf"
    assert sanitize_filename("../../../malicious.docx") == "malicious.docx"

def test_validator_rejects_executable_spoofing():
    # An executable PE (starts with MZ) renamed to .pdf
    fake_pdf = b"MZ\x90\x00\x03\x00\x00\x00" + b"\x00" * 100
    with pytest.raises(Exception) as exc_info:
        validate_uploaded_file("trojan.pdf", fake_pdf, "application/pdf")
    assert "توقيع" in str(exc_info.value) or "تنفيذي" in str(exc_info.value)

def test_validator_rejects_fake_docx():
    # Random text disguised as docx
    fake_docx = b"<html><script>alert(1)</script></html>"
    with pytest.raises(Exception) as exc_info:
        validate_uploaded_file("payload.docx", fake_docx, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    assert "توقيع" in str(exc_info.value) or "Word" in str(exc_info.value)

def test_validator_rejects_unsupported_extensions():
    with pytest.raises(Exception) as exc_info:
        validate_uploaded_file("script.py", b"print('hello')", "text/x-python")
    assert "غير مدعوم" in str(exc_info.value)

def test_validator_rejects_oversized_files():
    huge_bytes = b"0" * (50 * 1024 * 1024 + 10)
    with pytest.raises(Exception) as exc_info:
        validate_uploaded_file("huge.txt", huge_bytes, "text/plain")
    assert "يتجاوز الحد الأقصى" in str(exc_info.value)

# -------------------------------------------------------------
# 2. Integration Tests via API Upload
# -------------------------------------------------------------

@pytest.mark.asyncio
async def test_api_rejects_disguised_executable():
    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as ac:
        token = await get_or_create_user(ac, "security_tester@example.com", "فاحص الأمان")
        headers = {"Authorization": f"Bearer {token}"}

        fake_pdf = b"MZ" + b"\x00" * 1000
        res = await ac.post(
            "/api/v1/documents/upload",
            headers=headers,
            data={"title": "ملف تنفيذي مموه", "subject": "الفيزياء"},
            files={"file": ("malware.pdf", fake_pdf, "application/pdf")}
        )
        assert res.status_code == 400
        assert "توقيع" in res.json()["detail"] or "تنفيذي" in res.json()["detail"]

@pytest.mark.asyncio
async def test_api_handles_path_traversal_safely():
    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as ac:
        token = await get_or_create_user(ac, "traversal_tester@example.com", "فاحص المسارات")
        headers = {"Authorization": f"Bearer {token}"}

        sample_pdf_path = os.path.join(os.path.dirname(__file__), "..", "sample_physics_book.pdf")
        with open(sample_pdf_path, "rb") as f:
            pdf_bytes = f.read()

        res = await ac.post(
            "/api/v1/documents/upload",
            headers=headers,
            data={"title": "مستند مسار مفخخ", "subject": "الفيزياء"},
            files={"file": ("../../../../dangerous.pdf", pdf_bytes, "application/pdf")}
        )
        assert res.status_code == 201
        data = res.json()
        assert not data["filename"].startswith("..")
        assert "dangerous.pdf" in data["filename"]
        assert data["file_type"] == "pdf"

@pytest.mark.asyncio
async def test_docx_upload_and_chunk_extraction():
    # Build a valid Word (.docx) document in memory with headings, paragraphs, and tables
    doc = DocxDocument()
    doc.add_heading("قوانين الحركة لنيوتن في الفيزياء", level=1)
    doc.add_paragraph("ينص القانون الأول لنيوتن على أن الجسم الساكن يبقى ساكناً ما لم تؤثر عليه قوة خارجية.")
    doc.add_heading("جدول مقارنة القوانين", level=2)
    
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "القانون"
    table.cell(0, 1).text = "الصيغة"
    table.cell(1, 0).text = "القانون الثاني"
    table.cell(1, 1).text = "F = m * a"

    docx_stream = io.BytesIO()
    doc.save(docx_stream)
    docx_bytes = docx_stream.getvalue()

    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as ac:
        token = await get_or_create_user(ac, "docx_user@example.com", "مستخدم وورد")
        headers = {"Authorization": f"Bearer {token}"}

        res = await ac.post(
            "/api/v1/documents/upload",
            headers=headers,
            data={"title": "مذكرة فيزياء وورد", "subject": "الفيزياء"},
            files={"file": ("physics_notes.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        )
        assert res.status_code == 201
        doc_data = res.json()
        doc_id = doc_data["id"]
        assert doc_data["file_type"] == "docx"

        from app.services.document_worker import document_worker
        await document_worker.wait_for_document(doc_id, timeout=10.0)

        # Fetch chunks to verify source_type and content extraction
        chunks_res = await ac.get(f"/api/v1/documents/{doc_id}/chunks", headers=headers)
        assert chunks_res.status_code == 200
        chunks = chunks_res.json()
        assert len(chunks) > 0
        for chunk in chunks:
            assert chunk["source_type"] == "docx"
        
        all_chunk_text = " ".join(c["content"] for c in chunks)
        assert "نيوتن" in all_chunk_text
        assert "F = m * a" in all_chunk_text

@pytest.mark.asyncio
async def test_txt_upload_and_chunk_extraction():
    txt_content = (
        "مادة التاريخ - الحضارة المصرية القديمة\n"
        "تميزت الحضارة الفرعونية ببناء الأهرامات وتطور علوم الفلك والطب والزراعة على ضفاف نهر النيل.\n"
        * 10
    ).encode("utf-8")

    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as ac:
        token = await get_or_create_user(ac, "txt_user@example.com", "مستخدم نصوص")
        headers = {"Authorization": f"Bearer {token}"}

        res = await ac.post(
            "/api/v1/documents/upload",
            headers=headers,
            data={"title": "تاريخ الفراعنة", "subject": "التاريخ"},
            files={"file": ("history.txt", txt_content, "text/plain")}
        )
        assert res.status_code == 201
        doc_data = res.json()
        doc_id = doc_data["id"]
        assert doc_data["file_type"] == "txt"

        from app.services.document_worker import document_worker
        await document_worker.wait_for_document(doc_id, timeout=10.0)

        chunks_res = await ac.get(f"/api/v1/documents/{doc_id}/chunks", headers=headers)
        assert chunks_res.status_code == 200
        chunks = chunks_res.json()
        assert len(chunks) > 0
        for chunk in chunks:
            assert chunk["source_type"] == "txt"
        assert any("الأهرامات" in c["content"] for c in chunks)

@pytest.mark.asyncio
async def test_image_upload_and_ocr_routing():
    # Create a valid in-memory image
    img = Image.new("RGB", (300, 100), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)
    draw.text((20, 40), "StudyMind OCR Test", fill=(0, 0, 0))
    img_byte_arr = io.BytesIO()
    img.save(img_byte_arr, format="PNG")
    img_bytes = img_byte_arr.getvalue()

    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as ac:
        token = await get_or_create_user(ac, "image_user@example.com", "مستخدم صور")
        headers = {"Authorization": f"Bearer {token}"}

        res = await ac.post(
            "/api/v1/documents/upload",
            headers=headers,
            data={"title": "صورة ممسوحة ضوئياً", "subject": "الفيزياء"},
            files={"file": ("diagram_scan.png", img_bytes, "image/png")}
        )
        assert res.status_code == 201
        doc_data = res.json()
        doc_id = doc_data["id"]
        assert doc_data["file_type"] == "image"

        from app.services.document_worker import document_worker
        await document_worker.wait_for_document(doc_id, timeout=15.0)
        status_res = await ac.get(f"/api/v1/documents/{doc_id}/status", headers=headers)
        assert status_res.json()["status"].upper() == "READY"
